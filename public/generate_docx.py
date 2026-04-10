from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# Style helpers
def add_bottom_border(paragraph, color='00E5FF', width=3):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(width * 4))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_spacing(paragraph, before=0, after=0):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)

# === HEADER ===
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BXat')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0a, 0x0a, 0x0a)
set_spacing(p, after=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BaXiJen — acima de tudo, Brasileiro.')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x3D, 0x99, 0x42)
run.font.bold = True
set_spacing(p, after=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BaXiJen I.S. | baxijen.com.br')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
add_bottom_border(p, '00E5FF', 3)
set_spacing(p, after=12)

# === INTRO ===
p = doc.add_paragraph()
run = p.add_run('Imagine perguntar a um assistente virtual sobre uma norma municipal, um procedimento licitatório ou uma dúvida sobre FGTS, e receber uma resposta precisa, em português, com conhecimento real do contexto brasileiro. Não uma tradução genérica de uma IA americana. Uma resposta que entende o país onde você está.')
run.font.size = Pt(11.5)
set_spacing(p, after=4)

p = doc.add_paragraph()
run = p.add_run('É isso que o BXat faz.')
run.font.size = Pt(11.5)
run.font.bold = True
set_spacing(p, after=12)

# === O PROBLEMA ===
p = doc.add_paragraph()
run = p.add_run('O Problema')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0a, 0x0a, 0x0a)
add_bottom_border(p, 'E0E0E0', 2)
set_spacing(p, before=6, after=8)

p = doc.add_paragraph()
run = p.add_run('O Brasil está adotando inteligência artificial em ritmo acelerado. Organizações públicas e privadas buscam automação, eficiência e capacidade de resposta. Mas a maioria das soluções disponíveis no mercado carrega problemas estruturais que ninguém gosta de falar.')
run.font.size = Pt(11.5)
set_spacing(p, after=6)

# Dependência
p = doc.add_paragraph()
run = p.add_run('O primeiro é a dependência. ')
run.font.size = Pt(11.5)
run.font.bold = True
run2 = p.add_run('Quando uma organização contrata uma IA baseada em nuvens estrangeiras, seus dados — informações de cidadãos, estratégias institucionais, processos internos — trafegam para fora do país. Isso não é apenas uma questão de preferência. É uma violação potencial da Lei Geral de Proteção de Dados, que exige bases legais específicas para transferência internacional. Na prática, muitas organizações estão irregulares e não sabem.')
run2.font.size = Pt(11.5)
set_spacing(p, after=6)

# Qualidade
p = doc.add_paragraph()
run = p.add_run('O segundo é a qualidade. ')
run.font.size = Pt(11.5)
run.font.bold = True
run2 = p.add_run('Modelos globais foram treinados predominantemente em inglês, com dados e realidades de outros países. Eles sabem o que é um 401(k) americano, mas não sabem o que é um FGTS. Conhecem o regulamento de saúde dos Estados Unidos, mas não entendem como funciona o SUS. O resultado é uma IA que responde em português, mas pensa em inglês — e que frequentemente erra ou simplifica demais quando o assunto é Brasil. No BXat, o contexto brasileiro não é um adendo — é intrínseco ao modelo. FGTS, CLT, SUS, IPTU, licitações: tudo já faz parte da estrutura cognitiva da IA, não de uma camada superficial adicionada depois.')
run2.font.size = Pt(11.5)
set_spacing(p, after=6)

# Custo
p = doc.add_paragraph()
run = p.add_run('O terceiro é o custo. ')
run.font.size = Pt(11.5)
run.font.bold = True
run2 = p.add_run('Soluções de ponta cobram em dólar. Para um município brasileiro, uma assinatura de US$ 20 por usuário por mês se transforma em mais de R$ 115, sujeita a variação cambial. IA virou artigo de luxo quando deveria ser ferramenta de trabalho.')
run2.font.size = Pt(11.5)
set_spacing(p, after=12)

# === A SOLUÇÃO ===
p = doc.add_paragraph()
run = p.add_run('A Solução')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0a, 0x0a, 0x0a)
add_bottom_border(p, 'E0E0E0', 2)
set_spacing(p, before=6, after=8)

p = doc.add_paragraph()
run = p.add_run('O BXat é um assistente de inteligência artificial desenvolvido pela BaXiJen, uma startup brasileira nascida na pesquisa. Diferente das soluções importadas, o BXat foi pensado desde a origem para funcionar no Brasil, para brasileiros, respeitando as leis brasileiras.')
run.font.size = Pt(11.5)
set_spacing(p, after=8)

# Soberania
p = doc.add_paragraph()
run = p.add_run('Soberania de dados. ')
run.font.size = Pt(11.5)
run.font.bold = True
run.font.color.rgb = RGBColor(0x2D, 0x7A, 0x32)
run2 = p.add_run('O BXat opera em infraestrutura própria dentro do Brasil. Os dados não saem do país, não são armazenados em servidores estrangeiros, não alimentam modelos de terceiros. A conformidade com a LGPD não é uma adaptação posterior — é parte da arquitetura. Cada organização mantém controle total sobre onde seus dados estão, quem tem acesso e o que a IA pode fazer com eles.')
run2.font.size = Pt(11.5)
set_spacing(p, after=6)

# Inteligência
p = doc.add_paragraph()
run = p.add_run('Inteligência contextualizada. ')
run.font.size = Pt(11.5)
run.font.bold = True
run.font.color.rgb = RGBColor(0x2D, 0x7A, 0x32)
run2 = p.add_run('Construído sobre os modelos open-source mais avançados do mundo — como GLM-5 e Qwen 3.5 —, o BXat vai além do fine-tuning: atuamos diretamente na arquitetura da rede neural, otimizando a forma como o modelo processa e gera linguagem para o contexto brasileiro. O resultado é uma IA que não apenas foi treinada com dados nacionais — ela foi estruturalmente adaptada para pensar em português. Entende FGTS, CLT, INSS, IPTU, SUS, processos licitatórios e o ecossistema regulatório nacional. Quando um gestor municipal pergunta sobre uma normativa, o BXat responde com conhecimento de causa, não com uma aproximação traduzida de uma realidade que não é a nossa.')
run2.font.size = Pt(11.5)
set_spacing(p, after=6)

# Segurança
p = doc.add_paragraph()
run = p.add_run('Segurança em camadas. ')
run.font.size = Pt(11.5)
run.font.bold = True
run.font.color.rgb = RGBColor(0x2D, 0x7A, 0x32)
run2 = p.add_run('O BXat implementa guardrails de segurança baseados na tecnologia NeMo da NVIDIA — uma camada estrutural que filtra respostas inadequadas, impede a geração de conteúdo prejudicial e protege dados sensíveis de aparecerem em saídas. Não é um filtro cosmético. É um mecanismo que garante que a IA opere dentro dos limites definidos pela organização.')
run2.font.size = Pt(11.5)
set_spacing(p, after=6)

# Preço
p = doc.add_paragraph()
run = p.add_run('Preço em reais. ')
run.font.size = Pt(11.5)
run.font.bold = True
run.font.color.rgb = RGBColor(0x2D, 0x7A, 0x32)
run2 = p.add_run('Sem dependência de câmbio, sem surpresas na fatura. O modelo de assinatura é em moeda brasileira, com planos acessíveis para organizações de qualquer porte.')
run2.font.size = Pt(11.5)
set_spacing(p, after=12)

# === COMO FUNCIONA ===
p = doc.add_paragraph()
run = p.add_run('Como Funciona na Prática')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0a, 0x0a, 0x0a)
add_bottom_border(p, 'E0E0E0', 2)
set_spacing(p, before=6, after=8)

p = doc.add_paragraph()
run = p.add_run('A implantação é simples. O BXat é instalado na infraestrutura do cliente — servidor próprio, nuvem privada ou ambiente institucional. A organização define quem pode usar, o que a IA pode responder e quais dados estão acessíveis. A interface é uma conversa: o usuário pergunta, o BXat responde. Sem curva de aprendizado. Sem manual. Sem treinamento extensivo.')
run.font.size = Pt(11.5)
set_spacing(p, after=8)

for label, text in [
    ('Gestão pública municipal:', ' análise de dados urbanos, respostas sobre normativas, apoio ao atendimento ao cidadão, síntese de relatórios técnicos em linguagem acessível.'),
    ('Setor privado:', ' atendimento ao cliente, pesquisa, análise documental, suporte técnico — sem abrir mão da soberania dos dados.'),
    ('Instituições de ensino e pesquisa:', ' revisão bibliográfica, estruturação de projetos, interação com bases de dados acadêmicas, respeitando propriedade intelectual e privacidade.')
]:
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.font.size = Pt(11.5)
    run.font.bold = True
    run2 = p.add_run(text)
    run2.font.size = Pt(11.5)
    set_spacing(p, after=4)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# === A EQUIPE ===
p = doc.add_paragraph()
run = p.add_run('A Equipe')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0a, 0x0a, 0x0a)
add_bottom_border(p, 'E0E0E0', 2)
set_spacing(p, before=6, after=8)

p = doc.add_paragraph()
run = p.add_run('O BXat é desenvolvido pela BaXiJen I.S., uma startup fundada por três doutorandos em inteligência artificial com experiência direta em pesquisa aplicada e desenvolvimento de software. A equipe combina rigor acadêmico com agilidade de desenvolvimento, permitindo incorporar os mais recentes avanços em IA generativa sem depender de fornecedores estrangeiros.')
run.font.size = Pt(11.5)
set_spacing(p, after=6)

p = doc.add_paragraph()
run = p.add_run('A BaXiJen também opera com agentes de IA como força de trabalho interna — o que permite escalar a operação com eficiência e custo controlado, sem as despesas trabalhistas tradicionais que limitam startups brasileiras.')
run.font.size = Pt(11.5)
set_spacing(p, after=12)

# === O MOMENTO ===
p = doc.add_paragraph()
run = p.add_run('O Momento')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0a, 0x0a, 0x0a)
add_bottom_border(p, 'E0E0E0', 2)
set_spacing(p, before=6, after=8)

p = doc.add_paragraph()
run = p.add_run('A inteligência artificial está se tornando um fator determinante de competitividade e cidadania. Quando apenas grandes organizações com orçamentos em dólar têm acesso a IA de qualidade, a tecnologia deixa de ser ferramenta e passa a ser barreira — um mecanismo de exclusão que amplia desigualdades já estruturais. A democratização da IA não é um ideal romântico; é uma necessidade concreta. Municípios, pequenas empresas, instituições de ensino, organizações sociais — todos precisam de inteligência artificial que funcione no contexto deles, que custe o que eles podem pagar, que respeite as leis que os regem.')
run.font.size = Pt(11.5)
set_spacing(p, after=6)

p = doc.add_paragraph()
run = p.add_run('É isso que nos move. O BXat não existe para competir com soluções globais no topo do mercado. Existe para garantir que IA de qualidade não seja privilégio. Inteligência artificial brasileira, rodando em solo brasileiro, custando em moeda brasileira, entendendo problemas brasileiros — e acessível a quem precisa.')
run.font.size = Pt(11.5)
set_spacing(p, after=10)

# Closing
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Para nós.')
run.font.size = Pt(15)
run.font.bold = True
run.font.color.rgb = RGBColor(0x3D, 0x99, 0x42)
set_spacing(p, after=18)

# Top border for footer area
p = doc.add_paragraph()
add_bottom_border(p, '00E5FF', 2)
set_spacing(p, after=4)

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BaXiJen I.S.')
run.font.size = Pt(13)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0a, 0x0a, 0x0a)
set_spacing(p, after=0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('baxijen.com.br')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
set_spacing(p, after=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Documento informativo | Abril 2026')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.save('/home/leo-camilo/.openclaw/workspace/nextech/public/BXat-Apresentacao.docx')
print('DOCX saved!')